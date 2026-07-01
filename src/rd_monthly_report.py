"""Korea R&D intelligence monthly Word report for Fraunhofer Korea Office."""
from __future__ import annotations

import json
import logging
import os
from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from dotenv import load_dotenv
from openai import OpenAI

from src.config import PROJECT_ROOT, Settings
from src.daily_report import log_to_summarized_article, monthly_credibility_distribution
from src.rd_targeting import MONTHLY_RD_MIN_SCORE, compute_rd_match_score, parse_rd_fields, prepare_logs_for_monthly_rd

logger = logging.getLogger(__name__)

_HEADER_BLUE = RGBColor(0x1F, 0x39, 0x64)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = _HEADER_BLUE


def _compact_entry(log: dict, ref: int) -> dict:
    article = log_to_summarized_article(log)
    fields = parse_rd_fields(article.ko_summary_steps)
    return {
        "ref": ref,
        "date": log.get("log_date", ""),
        "title": log.get("title", ""),
        "url": log.get("url", ""),
        "source": log.get("source_name", ""),
        "score": log.get("rd_match_score") or compute_rd_match_score(article),
        "actor": fields.get("investment_actor", ""),
        "purpose": fields.get("investment_purpose", ""),
        "pain": fields.get("pain_point", ""),
        "strategy": fields.get("approach_strategy", ""),
        "proposable": article.rd_proposable_area,
        "fact": article.rd_fact_basis,
        "summary": article.ko_one_liner or article.llm_summary,
    }


def _synthesize_monthly_ko(year: int, month: int, entries: list[dict]) -> dict:
    load_dotenv(PROJECT_ROOT / ".env")
    client = OpenAI(
        api_key=os.getenv("OPENAI_API_KEY"),
        base_url=os.getenv("OPENAI_BASE_URL"),
    )
    prompt = f"""당신은 Fraunhofer 한국 사무소 R&D 전략가입니다.
{year}년 {month}월 국내 R&D 인텔리gence 항목 {len(entries)}건(적합도 {MONTHLY_RD_MIN_SCORE}점 이상)을 바탕으로 월간 보고서 JSON을 작성하세요.

규칙:
- 한국 국내 정부·기업 R&D 위탁·협력 기회만 다룸. 해외 시장·글로벌 벤더 분석 금지.
- 모든 문장 명사형 종결(-함/-임/-었음). -습니다/-합니다 금지.
- 소스에 없는 수치·기관명 추가 금지.

입력 데이터:
{json.dumps(entries, ensure_ascii=False)}

JSON 스키마:
{{
  "executive_summary": "3-4문장. 이번 달 국내 전력·에너지·ICT R&D 투자 트렌드",
  "opportunities": [
    {{"field": "분야명(HVDC/BESS/스마트그리드 등)", "summary": "정부·기업 움직임 요약", "refs": [1,2]}}
  ],
  "action_plan": [
    {{"target": "부처/기업명", "contact_angle": "접촉 논리", "rd_area": "제안 R&D", "refs": [1]}}
  ]
}}"""
    response = client.chat.completions.create(
        model=os.getenv("MODEL_NAME", "gemini-2.0-flash-lite"),
        response_format={"type": "json_object"},
        messages=[
            {
                "role": "system",
                "content": "국내 R&D 타겟팅 월간 보고서 작성. JSON만 반환.",
            },
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
    )
    raw = (response.choices[0].message.content or "{}").strip()
    return json.loads(raw)


def generate_rd_monthly_report(
    year: int,
    month: int,
    logs: list[dict],
    settings: Settings,
) -> Path:
    """Generate Korea-only R&D intelligence monthly Word report."""
    rd_logs, excluded = prepare_logs_for_monthly_rd(logs)
    if excluded:
        logger.info(
            "Excluded %d log(s) below R&D score %d for %04d-%02d",
            excluded,
            MONTHLY_RD_MIN_SCORE,
            year,
            month,
        )
    if not rd_logs:
        raise ValueError(
            f"No entries with R&D match score >= {MONTHLY_RD_MIN_SCORE} for {year}-{month:02d}."
        )

    compact = [_compact_entry(log, i) for i, log in enumerate(rd_logs, start=1)]
    try:
        structured = _synthesize_monthly_ko(year, month, compact)
    except Exception as exc:
        logger.warning("LLM monthly synthesis failed (%s) — using template fallback", exc)
        structured = _fallback_structure(compact)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"국내 R&D 인텔리전스 월간 보고서\n{year}년 {month}월")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = _HEADER_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Fraunhofer Institute Korea Office · Tech Market Intelligence Monitor")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x44, 0x54, 0x6A)

    doc.add_paragraph(
        f"분석 항목: {len(rd_logs)}건 (R&D 적합 {MONTHLY_RD_MIN_SCORE}점 이상) · "
        f"{monthly_credibility_distribution(rd_logs)}"
    )

    _add_heading(doc, "1. Executive Summary")
    doc.add_paragraph(structured.get("executive_summary", ""))

    _add_heading(doc, "2. Opportunities (분야별 R&D 기회)")
    for opp in structured.get("opportunities") or []:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{opp.get('field', '기타')}: ").bold = True
        p.add_run(opp.get("summary", ""))

    _add_heading(doc, "3. Action Plan (접촉 타겟)")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "타겟 (부처/기업)"
    hdr[1].text = "제안 R&D 영역"
    hdr[2].text = "접촉 논리"
    for action in structured.get("action_plan") or []:
        row = table.add_row().cells
        row[0].text = action.get("target", "")
        row[1].text = action.get("rd_area", "")
        row[2].text = action.get("contact_angle", "")

    _add_heading(doc, "4. 부록: 월간 R&D 스코어카드")
    score_table = doc.add_table(rows=1, cols=5)
    score_table.style = "Table Grid"
    sh = score_table.rows[0].cells
    sh[0].text = "점수"
    sh[1].text = "날짜"
    sh[2].text = "투자 주체"
    sh[3].text = "핵심 이슈"
    sh[4].text = "출처"
    for item in compact:
        row = score_table.add_row().cells
        row[0].text = f"{item['score']}/5"
        row[1].text = str(item["date"])
        row[2].text = item["actor"] or "—"
        row[3].text = (item["summary"] or item["title"])[:120]
        row[4].text = item["url"]

    output_dir = settings.reports_output_dir
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"rd-intelligence-report-{year}-{month:02d}-ko.docx"
    doc.save(output_path)
    logger.info("Generated R&D monthly report: %s", output_path)
    return output_path


def _fallback_structure(entries: list[dict]) -> dict:
    by_field: dict[str, list[dict]] = defaultdict(list)
    for e in entries:
        kw = e.get("proposable") or "기타"
        by_field[kw[:30]].append(e)

    opportunities = [
        {
            "field": field,
            "summary": f"{len(items)}건의 국내 R&D 신호 포착",
            "refs": [i["ref"] for i in items[:3]],
        }
        for field, items in list(by_field.items())[:5]
    ]
    action_plan = [
        {
            "target": e.get("actor") or e.get("source", ""),
            "contact_angle": e.get("strategy") or e.get("purpose", ""),
            "rd_area": e.get("proposable") or e.get("pain", ""),
            "refs": [e["ref"]],
        }
        for e in entries[:8]
        if e.get("actor")
    ]
    return {
        "executive_summary": (
            f"당월 R&D 적합 {MONTHLY_RD_MIN_SCORE}점 이상 항목 {len(entries)}건이 수집됨. "
            "정부·기업 투자 주체와 위탁 R&D 니즈를 스코어카드·Action Plan에 정리함."
        ),
        "opportunities": opportunities,
        "action_plan": action_plan,
    }
