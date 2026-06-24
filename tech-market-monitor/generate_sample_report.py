"""
Generate a sample May 2026 monthly report following the Fraunhofer TMR template
structure (colors, fonts, sections, table layouts).

Template reference: Fraunhofer_TMR_Template.docx
  - Cover: bold 26pt #1F4E79 title, 14pt #2E75B6 subtitle, bold 11pt #2E75B6 metadata
  - Heading 1: bold 14pt #1F4E79
  - Heading 2: bold 12pt #2E75B6
  - Body: 10pt black
  - Section tables: header row shaded #2E75B6 white bold text
"""

from __future__ import annotations

import calendar
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent
PROJECT_ROOT = BASE_DIR.parent
TEMPLATE_PATH = PROJECT_ROOT / "Fraunhofer_TMR_Template.docx"
REPORTS_DIR = PROJECT_ROOT / "output" / "monthly"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

YEAR, MONTH = 2026, 5
MONTH_NAME = calendar.month_name[MONTH]

# ── Brand colors (Fraunhofer TMR) ─────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x4E, 0x79)   # #1F4E79  — Heading 1, title
BLUE   = RGBColor(0x2E, 0x75, 0xB6)   # #2E75B6  — Heading 2, metadata
GRAY   = RGBColor(0x59, 0x59, 0x59)   # #595959  — placeholder body text
BLACK  = RGBColor(0x00, 0x00, 0x00)


# ── Report content (May 2026 sample data) ─────────────────────────────────────

TECH_NAME    = "AI Infrastructure & Smart Energy Technology"
REPORT_PERIOD = f"May {YEAR}"
PREPARED_BY   = "Tech Market Monitor — Automated Pipeline"
DEPARTMENT    = "Technology Intelligence Unit"
VERSION       = "v1.0 — Sample"
CLASSIFICATION = "Internal / Confidential"

TECHNOLOGY_SNAPSHOT = (
    "The convergence of AI-driven data center infrastructure and smart energy technology "
    "defines the most capital-intensive technology buildout of the decade. Hyperscalers are "
    "deploying modular, liquid-cooled compute facilities at gigawatt scale, while utilities "
    "race to upgrade grids capable of sustaining the associated power demand. Simultaneously, "
    "battery energy storage systems (BESS) and virtual power plants (VPP) are maturing from "
    "pilot projects into critical grid stability assets, catalyzed by AI-powered energy "
    "management systems that automate frequency regulation and demand response in real time."
)

KEY_FINDINGS = [
    "Market signal: Global AI data center capex commitments reached $120B (combined Microsoft, "
    "Google, Amazon, Meta) in Q1-Q2 2026, driving unprecedented demand for grid capacity upgrades. "
    "(Sources: Financial Times Tech, VentureBeat)",
    "Infrastructure signal: Grid-scale BESS deployments set a monthly record of 8 GWh commissioned "
    "globally in May 2026; AI-powered EMS is reducing renewable curtailment by up to 18%. "
    "(Source: VentureBeat)",
    "Korea-specific signal: Korean strategic investors (incl. BYD and domestic energy majors) "
    "stepped up grid-edge hardware stakes; MOTIE policy alignment with AI energy demand is expected "
    "to unlock additional BESS incentives in H2 2026. (Sources: The Verge, KOTRA estimates)",
    "Competitive signal: Tesla filed the 'Megapod' trademark for modular AI data center hardware, "
    "signaling a direct challenge to NVIDIA's 80%+ GPU market share in AI infrastructure. "
    "(Sources: MIT Technology Review, Electrek)",
    "Risk signal: HBM4 memory bandwidth has replaced GPU die as the critical supply constraint "
    "for AI training clusters; lead times of 18 months create a 2027 capacity cliff risk. "
    "(Source: Ars Technica)",
]

METRICS_ROWS = [
    ("Global AI Data Center Market", "$285B (2026E)", "+34% YoY", "$620B (2029E)", "Gartner / IDC"),
    ("Global BESS Deployment (May)", "8 GWh / month", "+210% YoY", "25 GWh / month", "VentureBeat / BloombergNEF"),
    ("Energy-Tech M&A (May 2026)", "$4.8B deal volume", "+180% YoY", "N/A", "Financial Times Tech"),
    ("TRL — Grid-AI Integration", "TRL 7-8", "Advancing", "TRL 9 by 2027", "Fraunhofer estimate"),
    ("Top Vendor Market Share (NVIDIA)", "~82% AI GPU", "Stable", "~70% (pressure)", "IDC MarketScape"),
]

TECH_DEFINITION = (
    "AI Infrastructure encompasses the hardware, software, and connectivity layers required to "
    "train, deploy, and serve large-scale AI models — including GPU/NPU clusters, high-bandwidth "
    "memory (HBM), InfiniBand/Ethernet fabrics, and the data center facilities housing them. "
    "Smart Energy Technology refers to the AI-augmented management of electrical grids, "
    "encompassing BESS, VPPs, demand response automation, and smart inverter/microgrid control "
    "systems. The two domains are converging as AI workload growth becomes the primary driver "
    "of global electricity demand growth, with IEA projecting data centers to consume 1,000+ TWh "
    "annually by 2026 — equivalent to Japan's total electricity consumption."
)

TECH_DIFF_ROWS = [
    ("GPU Compute (NVIDIA H200/B200)", "State-of-art AI training; HBM3e bandwidth 3.35 TB/s", "High", "TRL 9", "~$30-40K/unit"),
    ("Modular Data Center (Tesla Megapod)", "Self-contained AI compute unit; energy-efficient design", "Emerging", "TRL 5-6", "TBD"),
    ("Grid-Scale BESS (CATL / Tesla)", "4h+ discharge; AI-controlled frequency regulation", "High", "TRL 8-9", "$180-220/kWh"),
    ("VPP Platform (AutoGrid / Enel X)", "Aggregated DER dispatch; ML demand forecasting", "Moderate-High", "TRL 8", "SaaS model"),
    ("AI-EMS (Siemens / Honeywell)", "Real-time grid optimization; reduces curtailment 15-20%", "Moderate", "TRL 7-8", "Project-based"),
]

PATENT_ROWS = [
    ("EP4123456", "Modular liquid-cooled computing unit with AI workload scheduling", "Tesla Inc.", "2026", "US/EU", "Espacenet"),
    ("US18/234567", "Battery energy storage frequency regulation using reinforcement learning", "CATL", "2025", "US/CN", "Google Patents"),
    ("KR10-2025-0123456", "AI-based virtual power plant demand forecasting system", "KEPCO / KAIST", "2025", "KR", "KIPO"),
    ("EP4234567", "High-bandwidth memory controller for distributed AI inference", "SK Hynix", "2026", "EU/KR/US", "Espacenet"),
]

MARKET_OVERVIEW = (
    "The global AI infrastructure and smart energy technology market reached a combined estimated "
    "value of $285B in 2026, growing at a 5-year CAGR of approximately 28-34% (IDC, Gartner). "
    "May 2026 monitoring data indicates the market remains in a high-growth, high-investment phase: "
    "hyperscaler capex pledges of $120B for AI data centers are straining grid infrastructure in "
    "North America, Europe, and Northeast Asia. Utilities are responding with fast-tracked grid "
    "modernization programs, creating a secondary boom in BESS, smart inverters, and grid software. "
    "Energy-tech M&A volume hit $4.8B in May alone, with deal multiples averaging 12x ARR — "
    "signaling market confidence in long-term growth. (Sources: Financial Times Tech, VentureBeat, "
    "IDC H1 2026 Outlook)"
)

SEGMENTATION_ROWS = [
    ("Application", "AI Training / Inference", "$165B", "58%", "+36% CAGR", "GPU demand dominates"),
    ("Application", "Grid-Scale BESS", "$62B", "22%", "+31% CAGR", "Driven by renewable mandates"),
    ("Application", "VPP & Demand Response", "$28B", "10%", "+29% CAGR", "SaaS model growth"),
    ("End User", "Hyperscalers", "$178B", "63%", "+38% CAGR", "Concentrated buying power"),
    ("Deployment", "Cloud / Colocation", "$195B", "69%", "+35% CAGR", "On-prem declining share"),
    ("Region", "US / EU / Korea / Japan", "$260B", "91%", "+30% CAGR", "APAC fastest growing"),
]

REGIONAL_ROWS = [
    ("Korea", "$12B", "BESS mandate, AI national strategy", "MOTIE, MSIT, IITP", "KEPCO, Samsung SDI, LG Energy Solution", "KIET / KOTRA"),
    ("Japan", "$18B", "Data residency, energy security", "METI, NEDO", "SoftBank, NTT, Panasonic Energy", "METI"),
    ("China", "$55B", "AI sovereignty, CATL dominance", "MIIT, CAICT", "Huawei, CATL, BYD, Alibaba Cloud", "CAICT"),
    ("EU", "$48B", "AI Act compliance, REPowerEU", "EU Commission, BMBF", "Siemens Energy, Engie, Schneider", "Eurostat"),
    ("US", "$95B", "IRA incentives, hyperscaler CapEx", "NIST, DOE", "NVIDIA, Microsoft, Tesla, NextEra", "IDC"),
    ("SE Asia", "$8B", "Digital economy growth", "EDB Singapore, IMDA", "Google APAC, Temasek portfolio", "ADB"),
    ("India", "$6B", "PLI schemes, UPI data center boom", "NITI Aayog, DST", "Adani Green, Tata Power", "ADB"),
]

DRIVERS_BARRIERS = [
    ("Driver", "Hyperscaler AI CapEx surge ($120B, Q1-Q2 2026)", "High", "Financial Times Tech"),
    ("Driver", "Renewable energy integration mandates (EU REPowerEU, US IRA)", "High", "IEA / EU Commission"),
    ("Driver", "AI model size growth requiring 10x compute per generation", "High", "Gartner 2026 Hype Cycle"),
    ("Driver", "Korea BESS incentive programs under MOTIE 2026 roadmap", "Medium-High", "MOTIE / KIET"),
    ("Barrier", "HBM4 supply constraint — 18-month lead times (SK Hynix, Micron)", "High", "Ars Technica / IDC"),
    ("Barrier", "Grid infrastructure lag — utilities unable to fast-track approvals", "High", "Financial Times Tech"),
    ("Barrier", "EU AI Act compliance uncertainty delaying enterprise deployments", "Medium", "MIT Technology Review"),
    ("Barrier", "High capital intensity — BESS projects require $180-220/kWh capex", "Medium", "BloombergNEF"),
]

VENDOR_ROWS = [
    ("NVIDIA", "US", "Public", "H200/B200 GPUs, NIM microservices, DGX Cloud", "Full-stack AI platform lock-in", "Leader", "IDC MarketScape"),
    ("Microsoft Azure", "US", "Public", "Azure AI, OpenAI partnership, Megapod-ready DCs", "Vertical integration, OpenAI exclusivity", "Leader", "Gartner MQ"),
    ("Tesla Energy / Megapod", "US", "Public", "Megapod (AI DC HW, TM), Megapack BESS", "Energy efficiency differentiation", "Challenger", "Financial Times Tech"),
    ("CATL", "CN", "Public", "EnerOne BESS, CTP3 cells, VPP platform", "Cell manufacturing cost leadership", "Leader (BESS)", "IDC / KIET"),
    ("SK Hynix", "KR", "Public", "HBM3e/HBM4, AI DRAM", "HBM supply position for NVIDIA", "Critical supplier", "KOTRA / IDC"),
    ("ETRI / KAIST", "KR", "Government/Academic", "AI-EMS R&D, K-BESS standards", "Government-backed applied research", "Key R&D player", "KOTRA / KIET"),
]

KOREA_COMPETITIVE = (
    "Korea occupies a strategically critical position in both AI infrastructure and smart energy "
    "technology. SK Hynix is the primary supplier of HBM3e memory to NVIDIA, capturing an estimated "
    "50%+ of the HBM market (IDC). Samsung SDI and LG Energy Solution are top-3 global BESS cell "
    "suppliers. Government-backed research (ETRI, KAIST, POSTECH) is active in AI-EMS and VPP "
    "platforms, supported by IITP and KISTEP funding programs. KEPCO is piloting grid-scale AI "
    "optimization in partnership with domestic startups. Key risks: dependency on NVIDIA for AI "
    "compute demand, and potential US/EU export control escalation affecting chip supply chains. "
    "(Sources: KOTRA, KIET, MOTIE 2026 strategy documents)"
)

SWOT = {
    "Strengths": [
        "SK Hynix / Samsung SDI hold dominant HBM and BESS cell supply positions globally.",
        "ETRI and KAIST provide world-class applied R&D in AI-EMS and grid optimization.",
        "Strong government policy alignment (MOTIE, MSIT, IITP) with dedicated R&D budgets.",
    ],
    "Weaknesses": [
        "Limited domestic AI chip design capability; high dependency on NVIDIA architecture.",
        "KEPCO's grid modernization pace constrained by regulatory approval timelines.",
        "Korea-based VPP and demand-response software ecosystem is nascent vs. US/EU.",
    ],
    "Opportunities": [
        "Fraunhofer Korea can bridge EU-Korea standards collaboration for grid interoperability.",
        "Growing demand for K-BESS certification and testing services from export-oriented SMEs.",
        "IITP and Horizon Europe co-funding opportunities for AI-grid integration research.",
    ],
    "Threats": [
        "US CHIPS Act restrictions could limit Korean foundry access to advanced US IP.",
        "Chinese BESS cost leadership (CATL) threatening Korean cell manufacturers' margins.",
        "Rapid NVIDIA architectural shifts may erode SK Hynix HBM design partnerships.",
    ],
}

PORTER_ROWS = [
    ("Supplier Power", "High", "HBM supply concentrated in SK Hynix / Micron; TSMC dominates leading-edge fab", "Diversify memory suppliers; invest in CoWoS alternatives"),
    ("Buyer Power", "Medium", "Hyperscalers have significant bargaining power; enterprise buyers fragmented", "Long-term supply agreements reduce hyperscaler leverage"),
    ("Competitive Intensity", "Very High", "NVIDIA, Google TPU, AMD, Intel Gaudi competing in AI silicon; CATL vs LG in BESS", "Differentiation via energy efficiency and integration"),
    ("New Entrants", "Medium-High", "Tesla Megapod, Chinese OEMs entering AI HW; capital barriers moderate", "IP moats and standards participation critical"),
    ("Substitutes", "Low-Medium", "No near-term substitute for GPU-based AI training; pumped hydro vs BESS gradual", "Monitor solid-state battery commercialization timeline"),
]

ACADEMIC_ROWS = [
    ("Efficient Large-Scale AI Training with Sparse Mixture-of-Experts", "Fedus et al.", "NeurIPS 2025", "2025", "2,341", "doi:10.5555/3666122.3666234"),
    ("Grid Frequency Regulation via Deep Reinforcement Learning in BESS", "Kim J. et al.", "IEEE Trans. Smart Grid", "2025", "412", "doi:10.1109/TSG.2025.1234567"),
    ("Virtual Power Plant Aggregation with Federated Learning", "Park S. et al.", "ETRI Journal", "2026", "198", "doi:10.4218/etrij.2026-0034"),
    ("HBM4 Bandwidth Architecture for Next-Gen AI Accelerators", "SK Hynix / ISSCC", "ISSCC 2026", "2026", "87", "doi:10.1109/ISSCC.2026.9876543"),
    ("AI-Driven Demand Response Optimization in Smart Grids", "Chen X. et al.", "Applied Energy", "2026", "156", "doi:10.1016/j.apenergy.2026.01234"),
]

RD_FUNDING_ROWS = [
    ("K-BESS 2030 National Program", "MOTIE / IITP", "Korea", "KRW 850B", "AI-EMS, grid-scale BESS, VPP", "2024-2030", "motie.go.kr"),
    ("AI Semiconductor Initiative", "MSIT / IITP", "Korea", "KRW 1.2T", "AI chip design, HBM, NPU", "2023-2028", "iitp.kr"),
    ("Horizon Europe — Green Digital Transition", "EU Commission", "EU", "EUR 12B", "AI for energy, smart grids", "2021-2027", "ec.europa.eu"),
    ("BMBF — Energiesysteme der Zukunft", "BMBF", "Germany", "EUR 480M", "Grid digitalization, storage", "2024-2027", "bmbf.de"),
    ("DOE Grid Modernization Initiative", "NIST / DOE", "US", "$3.5B", "Grid resilience, AI integration", "2025-2030", "doe.gov"),
    ("NEDO Green Innovation Fund", "NEDO", "Japan", "JPY 2T", "Storage, offshore wind, AI-EMS", "2021-2030", "nedo.go.jp"),
]

EMERGING_RESEARCH = [
    "AI-native grid orchestration: integration of transformer-based forecasting models directly "
    "into BESS and VPP dispatch engines, enabling sub-second grid balancing decisions without "
    "human operator intervention (IEEE Xplore 2026 trend).",
    "Neuromorphic computing for edge-AI in smart meters and EV chargers: low-power inference "
    "chips embedded in grid endpoints, reducing latency in demand response by 80% vs. cloud-based "
    "approaches (arXiv cs.AI, Q1 2026 preprints).",
    "Hydrogen-battery hybrid storage for multi-day grid resilience: white-space R&D opportunity "
    "at the intersection of electrolyzer control AI and long-duration storage, directly relevant "
    "to Fraunhofer Korea's energy systems mandate (Fraunhofer ISE research direction, 2026).",
]

POLICY_ROWS = [
    ("EU AI Act — High-Risk AI Systems", "EU Commission", "EU", "Enforcement", "Aug 2026", "High", "ec.europa.eu"),
    ("EU Battery Regulation (2023/1542)", "EU Commission", "EU", "Active", "Feb 2027 (phase-in)", "High", "ec.europa.eu"),
    ("Korea BESS Safety Standards Revision", "MOTIE", "Korea", "Active", "Jun 2026", "High", "motie.go.kr"),
    ("US Executive Order on AI Energy Efficiency", "NIST / DOE", "US", "Active", "Apr 2026", "Medium-High", "nist.gov"),
    ("Korea AI Basic Act", "MSIT", "Korea", "Active", "Jan 2026", "Medium", "msit.go.kr"),
    ("OECD AI Principles — Revised 2025", "OECD", "Global", "Active", "2025", "Medium", "oecd-ilibrary.org"),
]

COMPLIANCE_ITEMS = [
    "Data governance: EU AI Act Article 10 requires high-risk AI systems (incl. grid-critical AI-EMS) "
    "to maintain data lineage, bias audit logs, and human override mechanisms. Korean counterpart "
    "requirements under the AI Basic Act are less prescriptive but evolving.",
    "Intellectual property: AI training on grid operational data may trigger IP obligations under "
    "KEPCO's data licensing terms. Standard-essential patent considerations apply to BESS frequency "
    "regulation protocols standardized under IEC 62933 series.",
    "Export control: Advanced AI chip exports (NVIDIA H100+, HBM3e) subject to BIS EAR controls; "
    "Fraunhofer Korea procurement of US-origin AI hardware requires end-user certificate compliance. "
    "CATL BESS cells face potential US Section 301 tariff risk for US-destination projects.",
]

HYPE_CYCLE = {
    "current_phase": "Slope of Enlightenment (AI Infrastructure) / Peak of Expectations (Grid-AI Integration)",
    "time_to_plateau": "2-3 years (AI Infrastructure); 4-5 years (Grid-AI Integration)",
    "source": "Gartner Hype Cycle for Artificial Intelligence 2026 / Gartner Hype Cycle for Utilities 2025",
}

PREDICTIONS_ROWS = [
    ("2027", "HBM4 supply normalizes; AI training cost per FLOP declines 40% YoY", "High", "NVIDIA margin compression; new entrants viable", "IDC"),
    ("2027", "EU mandates AI audit trails for grid-critical systems under AI Act Article 22", "High", "Compliance tooling market expands to $8B", "IDC / Gartner"),
    ("2028", "Grid-scale BESS reaches $150/kWh installed cost; crossover with peaker plants", "Medium-High", "Utility capex shifts from gas to BESS-dominant portfolios", "IDC / Gartner"),
    ("2029", "VPP platforms aggregate >200 GW in APAC; AI dispatch replaces manual grid ops", "Medium", "KEPCO-scale utilities become software-defined grid operators", "Gartner"),
    ("2030", "AI data center power demand exceeds 5% of global electricity; co-location with renewables standard", "High", "Energy tech and AI infra vertically integrated by hyperscalers", "McKinsey MGI"),
]

ROADMAP_ROWS = [
    ("Short-term", "0-2 years", "Megapod commercial launch; HBM4 mass production; K-BESS 2030 Phase 1 grants", "TRL 7-8", "HBM supply ramp; MOTIE grant approval", "HBM lead time risk"),
    ("Mid-term",   "2-4 years", "AI-EMS TRL 9 deployment; BESS $150/kWh milestone; EU Battery Regulation full enforcement", "TRL 8-9", "AI Act compliance tooling; IEC 62933 standard", "Regulatory delay risk"),
    ("Long-term",  "5-10 years","Software-defined grid mainstream; AI compute co-located with renewable generation", "TRL 9", "Policy: renewable mandate targets", "Technology consolidation"),
]

OPPORTUNITY_ROWS = [
    ("K-BESS Testing & Certification", "Research / Advisory", "MOTIE / KIET / Korean SMEs", "MOTIE / IITP", "High", "2026-2027"),
    ("EU-Korea Grid Interoperability Standard", "Research / Policy", "ETRI, Fraunhofer ISE, KEPCO", "BMBF / Horizon Europe", "High", "2026-2028"),
    ("AI-EMS Joint R&D Program", "Research", "KAIST / POSTECH / ETRI", "IITP / BMBF", "Medium-High", "2027"),
    ("Tech Transfer — VPP Platform to Korean SME", "Commercialization", "Korean energy SMEs (e.g., Encored)", "MOTIE / KISTEP", "Medium", "2027-2028"),
    ("Fraunhofer AI Act Compliance Advisory", "Advisory", "MSIT / Korean chaebol AI teams", "Consulting revenue", "Medium-High", "2026"),
]

RISK_ROWS = [
    ("US BIS export controls tighten on advanced AI chips", "Medium", "High", "Develop EU-sourced compute alternative; engage with NIST policy dialogue", "Technology Intelligence Unit"),
    ("CATL cost advantage erodes Korean BESS cell margin by 2027", "Medium-High", "Medium", "Accelerate differentiation via AI-EMS integration; target premium segments", "Energy Technology Team"),
    ("EU AI Act compliance delays Fraunhofer Korea's AI-EMS research deployment", "Low-Medium", "Medium", "Conduct internal Article 10 pre-compliance audit; engage EU liaison office", "Legal / Research"),
    ("IITP funding programme discontinuation (budget cycle risk)", "Low", "High", "Maintain dual funding pipeline: Horizon Europe + MOTIE + BMBF", "Strategy / Finance"),
]

RECOMMENDED_ACTIONS = [
    "Schedule a structured technology briefing with KISTEP and IITP (target: July 2026) to align "
    "on current TRL assessment methodology for AI-EMS platforms and validate Fraunhofer Korea's "
    "positioning against the K-BESS 2030 Phase 1 roadmap.",
    "Initiate a patent freedom-to-operate analysis via Espacenet and KIPO covering Tesla Megapod "
    "filings and SK Hynix HBM4 controller patents to identify white-space IP opportunities for "
    "Fraunhofer's modular grid-edge compute research.",
    "Draft a Horizon Europe partnership proposal (KA2 / EIC Pathfinder) targeting EU-Korea grid "
    "interoperability standards, engaging ETRI and Fraunhofer ISE as co-applicants; submit by "
    "September 2026 call deadline.",
]

METHODOLOGY = (
    "Data collection is performed via automated daily pipeline runs (08:00 KST) pulling from "
    "RSS feeds (MIT Technology Review, VentureBeat, Ars Technica, The Verge, Financial Times Tech), "
    "arXiv API (cs.AI, cs.LG, econ.GN categories), and Semantic Scholar. Items are filtered against "
    "a curated keyword list covering AI infrastructure, smart grid, BESS, and energy management "
    "system terms. Matched items are summarized via LLM (Gemini 2.0 Flash) with structured prompts "
    "ensuring business-impact focus. Monthly aggregation extracts trend signals across all logged "
    "items, producing this report via the Fraunhofer TMR template format."
)

DATA_QUALITY = (
    "Known limitations: (1) Korean SME revenue and funding data has limited public availability; "
    "KOTRA estimates used as proxies. (2) Chinese market data (CATL, Huawei) relies on CAICT "
    "publications which may have 6-12 month reporting lags. (3) Patent analysis is indicative only "
    "— full FTO assessment requires dedicated IP counsel review. (4) LLM summaries may miss nuance "
    "in highly technical academic abstracts; IEEE/ACM source articles should be consulted directly "
    "for scientific claims."
)

COLLECTED_ITEMS = [
    {
        "source": "MIT Technology Review",
        "title": "Inside the race to build liquid-cooled AI data centers at gigawatt scale",
        "summary": "Hyperscalers are investing billions in next-generation liquid-cooled facilities as AI GPU power densities exceed 100 kW per rack. Microsoft and Google are piloting direct liquid cooling at scale, targeting a 30% reduction in PUE.",
        "url": "https://www.technologyreview.com/2026/05/12/liquid-cooled-ai-datacenters/",
        "date": "2026-05-12",
        "keyword": "data center / AI infrastructure",
    },
    {
        "source": "MIT Technology Review",
        "title": "The EU AI Act is reshaping how enterprises buy software",
        "summary": "With enforcement deadlines approaching, vendors are bundling compliance dashboards and audit trails into enterprise AI platforms. Legal teams now sit alongside CTOs in AI procurement decisions.",
        "url": "https://www.technologyreview.com/2026/05/20/eu-ai-act-enterprise-software/",
        "date": "2026-05-20",
        "keyword": "AI governance / regulatory",
    },
    {
        "source": "VentureBeat",
        "title": "Grid-scale BESS hits record 8 GWh deployed in a single month",
        "summary": "Global battery energy storage deployments set a monthly record in May 2026, driven by renewable integration mandates in the US, EU, and South Korea. AI-powered EMS is reducing curtailment by up to 18%.",
        "url": "https://venturebeat.com/2026/05/08/bess-record-8gwh-monthly/",
        "date": "2026-05-08",
        "keyword": "BESS / battery energy storage",
    },
    {
        "source": "VentureBeat",
        "title": "Virtual power plants cross 50 GW milestone in Europe",
        "summary": "European VPP aggregators surpassed 50 GW of combined flexible capacity, making distributed energy resources a meaningful grid stability tool. Demand response automation is cutting peak load by 12% in pilot regions.",
        "url": "https://venturebeat.com/2026/05/15/vpp-50gw-europe/",
        "date": "2026-05-15",
        "keyword": "VPP / virtual power plant / demand response",
    },
    {
        "source": "VentureBeat",
        "title": "Energy-tech M&A: Three grid software acquisitions exceed $1B each in May",
        "summary": "Utilities and infrastructure funds closed three major acquisitions of grid management software firms, betting on AI-energy convergence. Deal multiples averaged 12x ARR.",
        "url": "https://venturebeat.com/2026/05/22/grid-software-ma-may2026/",
        "date": "2026-05-22",
        "keyword": "M&A / market size",
    },
    {
        "source": "Ars Technica",
        "title": "TSMC confirms Japan and Arizona fab expansions to feed AI chip demand",
        "summary": "TSMC announced accelerated timelines for its Kumamoto Phase 2 and Arizona N2 fabs, targeting volume production by late 2027. The expansions target AI accelerator die with advanced CoWoS packaging.",
        "url": "https://arstechnica.com/tech-policy/2026/05/tsmc-fab-expansion-ai/",
        "date": "2026-05-06",
        "keyword": "supply chain / AI infrastructure",
    },
    {
        "source": "Ars Technica",
        "title": "HBM supply crunch is the new bottleneck for AI training clusters",
        "summary": "With GPU die supply improving, HBM4 memory bandwidth has become the critical constraint for large-scale AI training. SK Hynix and Micron face 18-month lead times.",
        "url": "https://arstechnica.com/hardware/2026/05/hbm-supply-crunch-ai/",
        "date": "2026-05-18",
        "keyword": "supply chain / AI infrastructure",
    },
    {
        "source": "The Verge",
        "title": "Tesla and BYD deepen stakes in grid-edge hardware startups",
        "summary": "Both EV giants made strategic investments in microgrid controller and smart inverter startups in May, signaling a push to monetize energy infrastructure beyond vehicles.",
        "url": "https://www.theverge.com/2026/05/25/tesla-byd-grid-edge-investment/",
        "date": "2026-05-25",
        "keyword": "startup funding / smart grid",
    },
    {
        "source": "The Verge",
        "title": "Startup funding in demand response and microgrid controls hits $2.3B in May",
        "summary": "VC and corporate investors poured a record $2.3B into demand response automation and microgrid control startups, triple the year-ago figure.",
        "url": "https://www.theverge.com/2026/05/29/demand-response-microgrid-funding/",
        "date": "2026-05-29",
        "keyword": "startup funding / demand response / microgrid",
    },
    {
        "source": "Financial Times Tech",
        "title": "Hyperscalers pledge $120B in AI data center capex -- and utilities are scrambling",
        "summary": "Combined AI data center capex commitments from Microsoft, Google, Amazon, and Meta reached $120B in Q1-Q2 2026, straining power grids near major cloud regions.",
        "url": "https://www.ft.com/content/ai-datacenter-capex-utilities-2026",
        "date": "2026-05-03",
        "keyword": "data center / AI infrastructure / market size",
    },
    {
        "source": "Financial Times Tech",
        "title": "AI governance compliance costs reshape enterprise IT budgets",
        "summary": "Large enterprises are allocating 8-12% of AI project budgets to governance and compliance tooling following EU AI Act enforcement. Startups offering automated audit trails are seeing 5x ARR growth.",
        "url": "https://www.ft.com/content/ai-governance-compliance-costs-2026",
        "date": "2026-05-27",
        "keyword": "AI governance / regulatory / market size",
    },
]


# ── Helper functions ──────────────────────────────────────────────────────────

def _shade_cell(cell, hex_color: str) -> None:
    """Fill a table cell with a solid background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_col_widths(table, widths_cm: list[float]) -> None:
    """Set table column widths in cm."""
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_cm):
                cell.width = Cm(widths_cm[idx])


def _header_row(table, labels: list[str], col_widths: list[float] | None = None) -> None:
    """Write a header row with blue background and white bold text."""
    row = table.rows[0]
    for idx, label in enumerate(labels):
        cell = row.cells[idx]
        _shade_cell(cell, "2E75B6")
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
    if col_widths:
        _set_col_widths(table, col_widths)


def _data_row(table, row_idx: int, values: list[str], bold_first: bool = False) -> None:
    row = table.rows[row_idx]
    for idx, val in enumerate(values):
        cell = row.cells[idx]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(val)
        run.font.size = Pt(9)
        if bold_first and idx == 0:
            run.bold = True
        run.font.color.rgb = BLACK


def _add_table(doc: Document, headers: list[str], data_rows: list[list[str]],
               col_widths: list[float] | None = None, bold_first: bool = False) -> None:
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers, col_widths)
    for i, row_data in enumerate(data_rows, start=1):
        _data_row(table, i, [str(v) for v in row_data], bold_first=bold_first)
    doc.add_paragraph()


def _h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def _body(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    if italic:
        run.italic = True


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK


def _page_break(doc: Document) -> None:
    doc.add_page_break()


def _divider(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Build document ─────────────────────────────────────────────────────────────

def build_report() -> Path:
    doc = Document()

    # ── Page margins (matching Fraunhofer template approx) ────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ─────────────────────────────────────────────────────────────────────────
    # COVER PAGE
    # ─────────────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    # Main title
    p = doc.add_paragraph()
    run = p.add_run("Technology Market Research Report")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = NAVY

    # Institute subtitle
    p = doc.add_paragraph()
    run = p.add_run("Fraunhofer Institute | Korea Office")
    run.font.size = Pt(14)
    run.font.color.rgb = BLUE

    doc.add_paragraph()
    _divider(doc)
    doc.add_paragraph()

    # Metadata block
    for label, value in [
        ("Technology Name", TECH_NAME),
        ("Report Period",   REPORT_PERIOD),
        ("Prepared by",     PREPARED_BY),
        ("Department",      DEPARTMENT),
        ("Version",         VERSION),
        ("Classification",  CLASSIFICATION),
    ]:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}:   ")
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_label.font.color.rgb = BLUE
        run_val = p.add_run(value)
        run_val.font.size = Pt(11)
        run_val.font.color.rgb = BLUE

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 1 -- EXECUTIVE BRIEF
    # ─────────────────────────────────────────────────────────────────────────
    _h1(doc, "SECTION 1 -- EXECUTIVE BRIEF")

    _h2(doc, "1.1 Technology Snapshot")
    _body(doc, TECHNOLOGY_SNAPSHOT)

    _h2(doc, "1.2 Key Findings")
    for finding in KEY_FINDINGS:
        _bullet(doc, finding)

    _h2(doc, "1.3 Metrics Dashboard")
    _add_table(
        doc,
        headers=["Indicator", "Current Value", "YoY Change", "3-Year Forecast", "Source"],
        data_rows=[list(r) for r in METRICS_ROWS],
        col_widths=[5.5, 3.5, 3.0, 3.5, 3.5],
        bold_first=True,
    )

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 2 -- TECHNOLOGY PROFILE
    # ─────────────────────────────────────────────────────────────────────────
    _h1(doc, "SECTION 2 -- TECHNOLOGY PROFILE")

    _h2(doc, "2.1 Technology Definition & Principles")
    _body(doc, TECH_DEFINITION)

    _h2(doc, "2.2 Technology Readiness Level (TRL)")
    _add_table(
        doc,
        headers=["Assessment", "Result", "Basis / Evidence"],
        data_rows=[
            ["Current TRL", "TRL 7-8 (AI-EMS) / TRL 8-9 (BESS)", "MOTIE 2026 roadmap; Fraunhofer ISE assessment"],
            ["Target TRL", "TRL 9", "K-BESS 2030 Phase 1 product milestone"],
            ["Est. time to TRL 9", "2-3 years (AI-EMS); 1-2 years (BESS)", "Expert consensus; KISTEP forecasting reports"],
            ["Benchmark technology", "Google TPU v5 (AI) / CATL EnerOne (BESS)", "IDC MarketScape 2026; BloombergNEF BESS report"],
        ],
        col_widths=[5.0, 5.5, 8.5],
        bold_first=True,
    )

    _h2(doc, "2.3 Technology Differentiation")
    _add_table(
        doc,
        headers=["Technology", "Key Differentiator", "Maturity", "TRL", "Cost Benchmark"],
        data_rows=[list(r) for r in TECH_DIFF_ROWS],
        col_widths=[4.5, 6.0, 2.5, 2.0, 4.0],
    )

    _h2(doc, "2.4 Patent Landscape")
    _body(doc, "Sources: Espacenet (EPO), KIPO, DPMA, JPO, CNIPA, Google Patents")
    _add_table(
        doc,
        headers=["Patent No.", "Title", "Applicant", "Year", "Jurisdiction", "Source"],
        data_rows=[list(r) for r in PATENT_ROWS],
        col_widths=[3.0, 6.5, 3.5, 1.5, 2.5, 2.0],
    )

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 3 -- MARKET ANALYSIS
    # ─────────────────────────────────────────────────────────────────────────
    _h1(doc, "SECTION 3 -- MARKET ANALYSIS")

    _h2(doc, "3.1 Market Overview")
    _body(doc, MARKET_OVERVIEW)

    _h2(doc, "3.2 Market Segmentation")
    _add_table(
        doc,
        headers=["Segment Type", "Segment Name", "Market Size", "Share", "Growth Rate", "Notes"],
        data_rows=[list(r) for r in SEGMENTATION_ROWS],
        col_widths=[3.5, 4.0, 2.5, 2.0, 2.5, 4.5],
    )

    _h2(doc, "3.3 Regional Deep-Dive")
    _add_table(
        doc,
        headers=["Region", "Market Size", "Key Driver", "Policy Context", "Key Players", "Source"],
        data_rows=[list(r) for r in REGIONAL_ROWS],
        col_widths=[2.5, 2.5, 3.5, 3.5, 4.5, 2.5],
    )

    _h2(doc, "3.4 Market Drivers & Barriers")
    _add_table(
        doc,
        headers=["Type", "Driver / Barrier", "Impact", "Source"],
        data_rows=[list(r) for r in DRIVERS_BARRIERS],
        col_widths=[2.0, 8.5, 2.5, 6.0],
    )

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 4 -- COMPETITIVE LANDSCAPE
    # ─────────────────────────────────────────────────────────────────────────
    _h1(doc, "SECTION 4 -- COMPETITIVE LANDSCAPE")

    _h2(doc, "4.1 Vendor Mapping")
    _add_table(
        doc,
        headers=["Vendor", "HQ", "Type", "Core Offering", "Strategy", "Market Position", "Source"],
        data_rows=[list(r) for r in VENDOR_ROWS],
        col_widths=[3.5, 1.5, 2.0, 4.5, 4.0, 2.5, 3.5],
    )

    _h2(doc, "4.2 Korea-Specific Competitive Context")
    _body(doc, KOREA_COMPETITIVE)

    _h2(doc, "4.3 SWOT Analysis")
    swot_table = doc.add_table(rows=4, cols=2)
    swot_table.style = "Table Grid"
    labels = [("Strengths", "Weaknesses"), ("Opportunities", "Threats")]
    quadrant_data = [
        (SWOT["Strengths"], SWOT["Weaknesses"]),
        (SWOT["Opportunities"], SWOT["Threats"]),
    ]
    for row_idx in range(4):
        label_row = row_idx // 2
        is_header = (row_idx % 2 == 0)
        for col_idx in range(2):
            cell = swot_table.rows[row_idx].cells[col_idx]
            if is_header:
                _shade_cell(cell, "2E75B6")
                p = cell.paragraphs[0]
                p.clear()
                run = p.add_run(labels[label_row][col_idx])
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
            else:
                items = quadrant_data[label_row][col_idx]
                p = cell.paragraphs[0]
                p.clear()
                for item in items:
                    run = p.add_run(f"- {item}\n")
                    run.font.size = Pt(9)
                    run.font.color.rgb = BLACK
    doc.add_paragraph()

    _h2(doc, "4.4 Porter's Five Forces Summary")
    _add_table(
        doc,
        headers=["Force", "Intensity", "Key Driver", "Strategic Implication"],
        data_rows=[list(r) for r in PORTER_ROWS],
        col_widths=[3.5, 2.5, 6.0, 7.0],
        bold_first=True,
    )

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 5 -- INNOVATION & R&D LANDSCAPE
    # ─────────────────────────────────────────────────────────────────────────
    _h1(doc, "SECTION 5 -- INNOVATION & R&D LANDSCAPE")

    _h2(doc, "5.1 Academic Publication Trends")
    _body(doc, "Sources: IEEE Xplore, ACM Digital Library, arXiv, Springer/Elsevier, ETRI Journal, KISTI")
    _add_table(
        doc,
        headers=["Title", "Author(s)", "Journal / Conference", "Year", "Citations", "DOI"],
        data_rows=[list(r) for r in ACADEMIC_ROWS],
        col_widths=[6.0, 3.0, 3.5, 1.5, 2.0, 3.5],
    )

    _h2(doc, "5.2 R&D Funding Programs")
    _add_table(
        doc,
        headers=["Program", "Funding Body", "Region", "Budget", "Key Focus", "Period", "Source"],
        data_rows=[list(r) for r in RD_FUNDING_ROWS],
        col_widths=[3.5, 2.5, 2.0, 2.5, 3.5, 2.5, 2.5],
    )

    _h2(doc, "5.3 Emerging Research Directions")
    for item in EMERGING_RESEARCH:
        _bullet(doc, item)

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 6 -- REGULATORY & POLICY ENVIRONMENT
    # ─────────────────────────────────────────────────────────────────────────
    _h1(doc, "SECTION 6 -- REGULATORY & POLICY ENVIRONMENT")

    _h2(doc, "6.1 Policy Tracker")
    _add_table(
        doc,
        headers=["Policy / Regulation", "Issuing Body", "Region", "Status", "Effective Date", "Tech Impact", "Source"],
        data_rows=[list(r) for r in POLICY_ROWS],
        col_widths=[4.5, 3.0, 2.0, 2.0, 2.5, 2.5, 2.5],
    )

    _h2(doc, "6.2 Compliance Considerations")
    for item in COMPLIANCE_ITEMS:
        _bullet(doc, item)

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 7 -- TECHNOLOGY FORECAST
    # ─────────────────────────────────────────────────────────────────────────
    _h1(doc, "SECTION 7 -- TECHNOLOGY FORECAST")

    _h2(doc, "7.1 Gartner Hype Cycle Position")
    for label, value in [
        ("Current Phase", HYPE_CYCLE["current_phase"]),
        ("Time to Plateau", HYPE_CYCLE["time_to_plateau"]),
        ("Source", HYPE_CYCLE["source"]),
    ]:
        p = doc.add_paragraph()
        run_l = p.add_run(f"{label}: ")
        run_l.bold = True
        run_l.font.size = Pt(10)
        run_l.font.color.rgb = BLACK
        run_v = p.add_run(value)
        run_v.font.size = Pt(10)
        run_v.font.color.rgb = GRAY
        run_v.italic = True

    _h2(doc, "7.2 IDC-Style 5-Year Predictions")
    _add_table(
        doc,
        headers=["Year", "Prediction", "Confidence", "Implication", "Source"],
        data_rows=[list(r) for r in PREDICTIONS_ROWS],
        col_widths=[1.5, 6.5, 2.5, 5.5, 3.0],
        bold_first=True,
    )

    _h2(doc, "7.3 Technology Roadmap")
    _add_table(
        doc,
        headers=["Phase", "Timeframe", "Milestone", "TRL", "Key Enabler", "Risk"],
        data_rows=[list(r) for r in ROADMAP_ROWS],
        col_widths=[2.0, 2.0, 5.5, 1.5, 4.5, 3.5],
        bold_first=True,
    )

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 8 -- STRATEGIC IMPLICATIONS FOR FRAUNHOFER KOREA
    # ─────────────────────────────────────────────────────────────────────────
    _h1(doc, "SECTION 8 -- STRATEGIC IMPLICATIONS FOR FRAUNHOFER KOREA")

    _h2(doc, "8.1 Opportunity Assessment")
    _add_table(
        doc,
        headers=["Opportunity", "Type", "Potential Partners", "Funding Source", "Priority", "Timeline"],
        data_rows=[list(r) for r in OPPORTUNITY_ROWS],
        col_widths=[4.0, 3.0, 4.5, 3.5, 2.5, 2.5],
    )

    _h2(doc, "8.2 Risk Register")
    _add_table(
        doc,
        headers=["Risk", "Likelihood", "Impact", "Mitigation", "Owner"],
        data_rows=[list(r) for r in RISK_ROWS],
        col_widths=[5.0, 2.5, 2.0, 6.0, 3.5],
    )

    _h2(doc, "8.3 Recommended Actions (Next 90 Days)")
    for idx, action in enumerate(RECOMMENDED_ACTIONS, 1):
        _bullet(doc, f"Action {idx}: {action}")

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 9 -- SOURCES & METHODOLOGY
    # ─────────────────────────────────────────────────────────────────────────
    _h1(doc, "SECTION 9 -- SOURCES & METHODOLOGY")

    _h2(doc, "9.1 Source Registry — Collected Items (May 2026)")
    by_source: dict[str, list[dict]] = {}
    for item in COLLECTED_ITEMS:
        by_source.setdefault(item["source"], []).append(item)
    for source, items in sorted(by_source.items()):
        p = doc.add_paragraph()
        run = p.add_run(source)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BLUE
        for item in items:
            p2 = doc.add_paragraph(style="List Bullet")
            run2 = p2.add_run(f"{item['title']}  [{item['date']}] — Keyword: {item['keyword']}")
            run2.bold = True
            run2.font.size = Pt(9)
            run2.font.color.rgb = BLACK
            p3 = doc.add_paragraph()
            p3.paragraph_format.left_indent = Cm(1.0)
            run3 = p3.add_run(item["summary"])
            run3.font.size = Pt(9)
            run3.font.color.rgb = GRAY
            p4 = doc.add_paragraph()
            p4.paragraph_format.left_indent = Cm(1.0)
            run4 = p4.add_run(item["url"])
            run4.font.size = Pt(8)
            run4.font.color.rgb = BLUE

    _h2(doc, "9.2 Automated Monitoring Methodology")
    _body(doc, METHODOLOGY)

    _h2(doc, "9.3 Data Quality & Limitations")
    _body(doc, DATA_QUALITY)

    doc.add_paragraph()
    _divider(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"-- End of Report -- | {TECH_NAME} | {REPORT_PERIOD} | Fraunhofer Institute Korea Office")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    # ── Save ──────────────────────────────────────────────────────────────────
    out_path = REPORTS_DIR / f"report_{YEAR}_{MONTH:02d}.docx"
    doc.save(str(out_path))
    return out_path


if __name__ == "__main__":
    path = build_report()
    print(f"Report saved: {path}")
